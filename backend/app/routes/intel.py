import io
from datetime import datetime
from urllib.parse import quote
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from fastapi import APIRouter, HTTPException, Response, Depends
from typing import Optional, Literal
from sqlalchemy.orm import Session
from app.models import IntelListResponse, FavoriteToggleRequest, ExportRequest, IntelItem, Tag
from app.database import get_db
from app import crud
from app.agent.orchestrator import orchestrator

router = APIRouter()

@router.get("/", response_model=IntelListResponse)
async def get_intel(
    type: Literal["hot", "history", "all"] = "all",
    q: Optional[str] = None,
    range: Literal["all", "3h", "6h", "12h"] = "all",
    limit: int = 20,
    offset: int = 0,
    db: Session = Depends(get_db)
):
    items, total = crud.get_filtered_intel(db, type_filter=type, q=q, range_filter=range, limit=limit, offset=offset)
    return {"items": items, "total": total}

@router.get("/favorites", response_model=IntelListResponse)
async def get_favorites(
    q: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
    db: Session = Depends(get_db)
):
    items, total = crud.get_favorites(db, q=q, limit=limit, offset=offset)
    return {"items": items, "total": total}

@router.post("/export")
async def export_intel(req: ExportRequest, db: Session = Depends(get_db)):
    # Determine items to export
    items = []
    
    if req.ids and len(req.ids) > 0:
        db_items = crud.get_by_ids(db, req.ids)

        by_id = {x.id: x for x in db_items}

        missing_ids = [x for x in req.ids if x not in by_id]
        cached_items = {}
        for missing_id in missing_ids:
            cached = orchestrator.get_cached_intel(missing_id)
            if not cached:
                continue

            tags = []
            for t in cached.get("tags") or []:
                if isinstance(t, dict):
                    tags.append(Tag(label=t.get("label", ""), color=t.get("color", "blue")))
                else:
                    tags.append(Tag(label=str(t), color="blue"))

            intel_item = IntelItem(
                id=str(cached.get("id", missing_id)),
                title=cached.get("title") or "",
                summary=cached.get("summary") or "",
                source=cached.get("source") or "Hot Stream",
                url=cached.get("url"),
                time=cached.get("time") or "",
                timestamp=float(cached.get("timestamp") or datetime.now().timestamp()),
                tags=tags,
                favorited=bool(cached.get("favorited") or False),
                is_hot=False,
                content=cached.get("content"),
                thing_id=cached.get("thing_id") or cached.get("thingId"),
            )

            if not crud.get_intel_by_id(db, intel_item.id):
                crud.create_intel_item(db, intel_item)

            cached_items[intel_item.id] = intel_item

        resolved_items = []
        unresolved_ids = []
        for requested_id in req.ids:
            if requested_id in by_id:
                resolved_items.append(by_id[requested_id])
                continue
            if requested_id in cached_items:
                resolved_items.append(cached_items[requested_id])
                continue
            unresolved_ids.append(requested_id)

        if unresolved_ids:
            raise HTTPException(status_code=404, detail=f"Items not found: {', '.join(unresolved_ids)}")

        items = resolved_items
    else:
        items, _ = crud.get_filtered_intel(
            db,
            type_filter=req.type or "all",
            q=req.q,
            range_filter=req.range or "all",
            limit=1000
        )
    
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(12)

    def _safe_text(v: Optional[str]) -> str:
        return (v or "").strip()

    def _add_kv_line(key: str, value: str):
        p = doc.add_paragraph()
        rk = p.add_run(f"{key}：")
        rk.bold = True
        p.add_run(value)
        return p

    def _add_center_title(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        return p

    def _add_body(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        return p

    for idx, item in enumerate(items):
        tags_value = " / ".join([t.label for t in (item.tags or []) if _safe_text(t.label)])
        if not tags_value:
            tags_value = "暂无"

        _add_kv_line("拟投栏目", tags_value)
        _add_kv_line("事件时间", _safe_text(item.time) or "暂无")
        _add_kv_line("价值点", _safe_text(item.summary) or "暂无")

        doc.add_paragraph()

        _add_center_title(_safe_text(item.title) or "未命名")

        body_text = _safe_text(item.content) or _safe_text(item.summary)
        if body_text:
            _add_body(body_text)

        source_parts = [f"来源：{_safe_text(item.source) or 'Unknown'}", f"原标题：{_safe_text(item.title) or '未命名'}"]
        if _safe_text(item.url):
            source_parts.append(f"来源URL：{_safe_text(item.url)}")
        doc.add_paragraph(f"（{'，'.join(source_parts)}）")

        if idx != len(items) - 1:
            doc.add_page_break()
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Determine filename
    filename = "情报批量导出.docx"
    if len(items) == 1:
        # Remove invalid chars for filename
        safe_title = "".join([c for c in items[0].title if c not in r'\/:*?"<>|'])
        filename = f"{safe_title}.docx"
    
    # Encode filename for header
    encoded_filename = quote(filename)
    
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )

@router.get("/{id}", response_model=IntelItem)
async def get_intel_detail(id: str, db: Session = Depends(get_db)):
    item = crud.get_intel_by_id(db, id)
    if not item:
        cached = orchestrator.get_cached_intel(id)
        if not cached:
            raise HTTPException(status_code=404, detail="Intel item not found")

        tags = []
        for t in cached.get("tags") or []:
            if isinstance(t, dict):
                tags.append(Tag(label=t.get("label", ""), color=t.get("color", "blue")))
            else:
                tags.append(Tag(label=str(t), color="blue"))

        intel_item = IntelItem(
            id=str(cached.get("id", id)),
            title=cached.get("title") or "",
            summary=cached.get("summary") or "",
            source=cached.get("source") or "Hot Stream",
            url=cached.get("url"),
            time=cached.get("time") or "",
            timestamp=float(cached.get("timestamp") or datetime.now().timestamp()),
            tags=tags,
            favorited=bool(cached.get("favorited") or False),
            is_hot=False,
            content=cached.get("content"),
            thing_id=cached.get("thing_id") or cached.get("thingId")
        )

        if not crud.get_intel_by_id(db, intel_item.id):
            crud.create_intel_item(db, intel_item)

        return intel_item
    
    tags = []
    for t in item.tags or []:
        if isinstance(t, dict):
            tags.append(Tag(label=str(t.get("label", "")), color=str(t.get("color", "blue"))))
        else:
            tags.append(Tag(label=str(t), color="blue"))
    return IntelItem(
        id=item.id,
        title=item.title,
        summary=item.summary,
        source=item.source,
        url=item.url,
        time=item.publish_time_str,
        timestamp=item.timestamp,
        tags=tags,
        favorited=item.favorited,
        is_hot=item.is_hot,
        content=item.content,
        thing_id=item.thing_id
    )

@router.post("/{id}/favorite")
async def toggle_favorite(id: str, req: FavoriteToggleRequest, db: Session = Depends(get_db)):
    item = crud.toggle_favorite(db, id, req.favorited)
    if item:
        return item

    cached = orchestrator.get_cached_intel(id)
    if not cached:
        raise HTTPException(status_code=404, detail="Item not found")

    tags = []
    for t in cached.get("tags") or []:
        if isinstance(t, dict):
            tags.append(Tag(label=t.get("label", ""), color=t.get("color", "blue")))
        else:
            tags.append(Tag(label=str(t), color="blue"))

    intel_item = IntelItem(
        id=str(cached.get("id", id)),
        title=cached.get("title") or "",
        summary=cached.get("summary") or "",
        source=cached.get("source") or "Hot Stream",
        url=cached.get("url"),
        time=cached.get("time") or "",
        timestamp=float(cached.get("timestamp") or datetime.now().timestamp()),
        tags=tags,
        favorited=bool(req.favorited),
        is_hot=bool(cached.get("is_hot") if cached.get("is_hot") is not None else True),
        content=cached.get("content"),
        thing_id=cached.get("thing_id") or cached.get("thingId"),
    )

    if not crud.get_intel_by_id(db, intel_item.id):
        crud.create_intel_item(db, intel_item)

    item = crud.toggle_favorite(db, intel_item.id, req.favorited)
    if not item:
        raise HTTPException(status_code=500, detail="Failed to toggle favorite")
    return item
